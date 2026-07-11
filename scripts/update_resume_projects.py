from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\28235\Desktop\简历.docx")
OUTPUT = Path(r"C:\Users\28235\Desktop\简历_项目经历更新版.docx")


PET_SHOP_PROJECT = """2026.06 - 2026.07    宠物商城 B/S 平台（Pet_shop）
项目内容：面向宠物用品购买、活体宠物浏览、宠物档案管理和商家商品管理等场景，开发一套前后端分离宠物商城系统，支持短信验证码登录、角色权限、商品审核、购物车、订单流转、宠物成长记录和 AI 导购推荐。
技术架构：后端 FastAPI + Pydantic + SQLAlchemy + Alembic + MySQL/PyMySQL；前端 uni-app / Vue H5；JWT 鉴权、RBAC 权限控制、阿里云短信 SDK、DeepSeek API；分层采用 Router / Service / Repository / Model / Schema。
基于 FastAPI 拆分认证、用户角色、商品、购物车、订单、宠物档案、活体宠物、平台审核和导购 Agent 等接口模块，统一封装响应、异常处理、依赖注入和数据库会话。
设计商品与订单业务流程，实现商品创建、编辑、提交审核、上架/下架、加入购物车、创建订单、模拟支付、取消订单、商家发货和确认收货等核心能力。
实现宠物档案与成长记录模块，支持新增/编辑宠物、设置当前宠物、记录健康与提醒信息，并在活体宠物购买后生成对应宠物档案。
开发购物导购 Agent，结合用户需求、预算、当前宠物档案和商品库存进行规则排序，并调用 DeepSeek 生成中文推荐总结，提升商品推荐的业务贴合度。
在前端使用 uni-app 页面路由组织账号、首页、商品详情、购物车、订单、宠物档案、商家商品和平台审核页面，完成 B/S 风格适配与权限按钮控制。"""


VEROXA_PROJECT = """2026.07 - 2026.07    企业级 AI 经营平台 MVP（Veroxa Enterprise）
项目内容：面向企业经营数据接入、风险分析、决策审批、策略评估、模拟执行和通知反馈场景，开发一套 MVP Web 平台，跑通“企业登录 -> Mock ERP 同步 -> 数据标准化 -> 规则分析 -> 决策审批 -> 策略评估 -> 执行任务 -> 通知 -> Today 聚合”的闭环。
技术架构：前端 Vue 3 + TypeScript + Vite + Pinia + Vue Router + Axios + Element Plus；后端 FastAPI + SQLAlchemy + Alembic + PostgreSQL + Redis + RQ Worker；JWT 认证、RBAC 权限、pytest / pytest-asyncio / httpx 自动化测试。
搭建企业注册登录、密码哈希、JWT 鉴权、企业级 RBAC 和操作日志能力，以 enterprise_id 作为多租户数据隔离边界，确保跨企业资源不可读取或操作。
实现 Connect、Data、AI Analyze、Decision、Policy、Execution、Notification、Today 等业务模块，按 routes / services / schemas / models 分层组织接口、业务编排和数据校验。
设计 Mock ERP Connector 与数据同步流程，将外部商品、库存、订单、采购单等原始数据标准化落库，并通过幂等边界避免重复同步。
建设规则分析、决策生成与审批、策略 allow / require_approval / block 三分支评估、模拟执行失败重试和 mock 通知能力，形成可验收的经营闭环。
完成 PostgreSQL 多 schema 数据库治理，将业务表按 core、connect、data、ai、decision、policy、execution、notification 等领域拆分，并使用 Alembic 管理可升级、可回滚迁移。
补充后端自动化验收测试，覆盖主路径闭环、异常分支、权限拒绝、跨企业隔离、幂等、策略分支、执行重试、通知幂等和 Today 聚合反馈。"""


def replace_project_section(text: str) -> str:
    marker = "\n\n\n\n自我评价"
    if marker not in text:
        raise RuntimeError("Could not find the self-evaluation marker in the resume.")
    insert = f"\n\n{PET_SHOP_PROJECT}\n\n{VEROXA_PROJECT}"
    return text.replace(marker, insert + marker, 1)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    lines = text.split("\n")
    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = paragraph.add_run(line)
        run.font.name = "宋体"


def main() -> None:
    doc = Document(SOURCE)
    target_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "项目经验" in cell.text and "AI 智能起名" in cell.text and "自我评价" in cell.text:
                    if not target_cells or cell is not target_cells[-1]:
                        target_cells.append(cell)

    if not target_cells:
        raise RuntimeError("Could not find the project-experience cell.")

    updated = replace_project_section(target_cells[0].text)
    for cell in target_cells:
        set_cell_text(cell, updated)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
